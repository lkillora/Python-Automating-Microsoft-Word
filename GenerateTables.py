import pandas as pd
import datetime
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import glob

filelist = []
for file_name in glob.glob('./'+'*.csv'):
    filelist += [(pd.read_csv(file_name), file_name)]
    
dates = [datetime.date(2020, 8, 31)]
for i in range(39):
    dates += [dates[-1] + datetime.timedelta(days=7)]
nonschoolweeks = [datetime.date(2019, 10, 26), 
                  datetime.date(2019, 12, 28), 
                  datetime.date(2020, 1, 4),
                 datetime.date(2020, 2, 15), 
                 datetime.date(2020, 3, 29),
                 datetime.date(2020, 4, 5)]
datenames = [d.strftime("%b") + " " + str(int(d.strftime("%d"))) for d in dates]

doc = docx.Document()
section = doc.sections[-1]

new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
for tuples in filelist:
    file = tuples[0]
    file_name = tuples[1][2:-4]
    file.sort_values("Full Name", inplace=True)
    num_groups = file.shape[0]//3 + 1 if file.shape[0]%3 > 0 else file.shape[0]//3
    groups = [[str(g+1)] for g in range(num_groups)]
    for r in range(file.shape[0]):
        groups[r//3] += [file["Full Name"].values[r]]
    
    if len(groups[-1]) == 2:
        groups[0] += [groups[-1][1]]
        groups = groups[:-1]
        num_groups -= 1
        for g in range(len(groups)):
            groups[g] += [""]*(5-len(groups[g]))
        names = ["Name " + str(i) for i in range(1, 5)]
    else:
        groups[-1] += [""]*(4-len(groups[-1]))
        names = ["Name " + str(i) for i in range(1, 4)]
    
    for d in range(len(datenames)):
        groups[d%num_groups] += [datenames[d]]
    num_weeks = len(dates)//num_groups + 1 if len(dates)%num_groups > 0 else len(dates)//num_groups
    df = pd.DataFrame(groups, columns=["Group"] + names +["Week "+str(w+1) for w in range(num_weeks)])
    #df.drop(columns="Group", inplace=True)
    df.fillna("", inplace=True)
    doc.add_heading('Class '+ file_name + " Cleaning Rota", 0)
    table = doc.add_table(rows=1, cols=df.shape[1])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells

    for c in range(df.shape[1]):
        hdr_cells[c].text = df.columns[c]

    for rw in range(0, df.shape[0]):
        row_cells = table.add_row().cells

        for c in range(df.shape[1]):
            if c == 0:
                width = Inches(0.05)
            if c >= 1 and c <= 3:
                width = Inches(0.5*5)
            else:
                width = Inches(0.1) 
            row_cells[c].width = width
            row_cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[c].text = df.values[rw][c]
    table.style = 'LightShading-Accent1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()
doc.save('./Rotas.docx')


